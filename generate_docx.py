# -*- coding: utf-8 -*-
"""
WPS 兼容 Word 文档：材料定义与计价规则 — 结构化改造方案
企业内用风格：精炼、问题直接可见、快速可扫描
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT = r"D:\desktop\材料定义与计价规则_结构化改造方案_v5.docx"

doc = Document()

# ── 页面设置 ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# ── 样式 ──
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(18)
    run.font.bold = True
    return p


def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    return p


def H1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(13)
    run.font.bold = True
    return p


def H2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(11.5)
    run.font.bold = True
    return p


def P(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.7)
    p.paragraph_format.line_spacing = 1.6
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(11)
    return p


def B(text):
    """要点句 —— 无缩进，问题直接可见"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.6
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(11)
    return p


def Q(text):
    """关键结论"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run.font.italic = True
    return p


def add_simple_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        '<w:tblPr %s></w:tblPr>' % nsdecls("w")
    )
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>' % nsdecls("w")
    )
    tblPr.append(borders)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(9.5)
        run.font.bold = True
        shading = parse_xml(
            '<w:shd %s w:fill="E8E8E8" w:val="clear"/>' % nsdecls("w")
        )
        cell._tc.get_or_add_tcPr().append(shading)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(9.5)
            if r % 2 == 1:
                shading = parse_xml(
                    '<w:shd %s w:fill="F5F5F5" w:val="clear"/>' % nsdecls("w")
                )
                cell._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph()
    return table


# ═══════════════════════════════════════
# 正文
# ═══════════════════════════════════════

add_title('材料定义与计价规则')
add_subtitle('结构化问题分析与改造方案')

# ── 一、问题定位 ──
H1('一、当前文档的问题定位')

P(
    '经审查，这份 Excel 文档共 56 行、25 列，单 Sheet，覆盖了材料主数据、'
    '计价公式、产品物料清单三项内容。文档在计价体系层面的逻辑是完整的——'
    '分类层级清晰，损耗率按材料大类分档而非逐行随意取值，公式结构在不同材料间保持统一。'
    '问题不在于内容本身的正误，而在于结构失配：'
    '一张平面表承载了超出其表达能力的信息密度，导致日常使用和维护场景中出现多处可规避的障碍。'
)

P('以下两个核心问题在企业使用场景中直接影响工作效率和数据可靠性。')

# ── 漏洞一 ──
H1('二、问题一：表格过于复杂，一张表混入三种互不兼容的职能')

H2('问题表现')

P(
    '这张表同时承担了三项职能——登记材料属性、定义计算公式、标注产品与材料的对应关系。'
    '三项职能的行粒度不同（材料按 SKU 一行，公式按规则一行，产品按关联一行），'
    '列需求不同（材料需要规格列和单价列，公式不需要；公式需要大段文本列，材料不需要），'
    '却被纳入同一套 25 列的列结构中。审查发现，没有任何一种使用场景需要用到全部 25 列：'
    '采购经理仅需前 10 列，报价工程师仅需后 3 列，'
    '项目经理需要的信息散落在 G 列的自由文本和各行备注之间。'
    '每一列对至少一种使用场景是冗余的，使用者需要在阅读过程中自行判断哪些列与当前任务相关。'
)

H2('典型后果')

B('1. 隐式继承导致增删改查极易出错。')
P(
    '表中同一材料的多个变体（如普通多层板 E0 12mm / E0 15mm / E1 12mm / E1 15mm）'
    '仅在首行填写名称、等级、说明等列，后续变体行留空，依赖阅读者向上追溯补全。'
    '这种「后续行留空表示继承上行」的写法使表中存在大量合并单元格，'
    '后续无论是增、删、改、查还是其他常规操作，都极易因继承链断裂而报出无法执行的错误，'
    '大幅增加了维护难度。'
)

B('2. 关键定义被多处复制，修订时无法保证一致。')
P(
    '「投影面积」的完整定义（约 250 字）在普通多层板、层板、木皮、封闭漆、辅材、加工费六处逐字重复。'
    '三种产品类型的用量公式在四个材料段落中各抄一遍。'
    '若需修订定义（如调整异形处理规则），须在六处同步修改，漏改一处即造成同一规则在不同位置给出不同表述，'
    '计价口径不再一致。'
)

B('3. 表头分类失效。')
P(
    'C1 标题「材料分类及其与使用场景偏好」管辖 C-G 列，实际混杂了标识、描述、分类、偏好、关联五种不同性质的字段。'
    'H1 标题「计算规则及其公式」管辖 H-P 列，却将计量单位、规格、单价等材料属性也纳入其中。'
    '表头已无法提供有效的列定位和导航。'
)

# ── 漏洞二 ──
H1('三、问题二：有效信息被参考数据淹没')

H2('问题表现')

P(
    '全表 56 行中，约 70% 的行标注了「本项目不使用」或「本项目内无此材料」。'
    '阻燃多层板、密度板、三聚氰胺板、欧松板、指接板、实木整板——这些材料的属性数据填写完整，'
    '但完全不参与本项目计价。审查确认，真正参与本项目的仅有 7 行：'
    'E0 普通多层板 12mm / 9mm、WD01 木皮、封边条、封闭漆、辅材、加工费。'
)

P(
    '参考材料本身具有一定保留价值（未来项目可能启用），但与活跃数据置于同一平面空间，'
    '仅依靠 G 列的一行文字标注区分。日常高频操作——查一个单价、确认一个损耗率——'
    '每次都需要在 56 行中定位目标行，反复过滤无用数据。'
    '此外，表格末尾 8 个成本类别（管理费用、运输费用、包装材料等）仅有标题行，其余 24 列全部空白；'
    'B 列整列空置。这些空白行列进一步增加了表格的视觉体积，却不提供任何有效信息。'
)

H2('直接影响')

B('- 无法快速统计本项目涉及的材料数量——须逐行检查 G 列，人工计数，易遗漏且难以复核。')
B('- 增删材料时定位模糊——新增材料是放入参考区还是活跃区，缺乏明确规则；参考材料转为启用时，'
   '是原地修改标注还是整体迁移至活跃区，未有约定。')
B('- 打印或分享时携带 40 余行无关数据，而实际需要的仅约 7 行。')

# ── 根因 ──
H1('四、两个问题的共同根因')

P(
    '审查表明，上述两个问题指向同一个结构性缺陷：'
    '这张表用空间关系（行序、合并单元格、纵向对齐）来表达业务逻辑，而非依赖字段值。'
    '「这些行同属板材」的信息由 A 列合并单元格的视觉范围暗示，而非在每一行的分类列中显式标注。'
    '「本行与上行材料相同」的信息由留空表达，而非每行各自填满。'
    '「本行不参与本项目」的信息由 G 列的自由文本承载，而非可筛选、可统计的状态字段。'
)

Q(
    '核心原则：信息应存放在字段值中，而非依赖于空间关系。'
    '排序、筛选、导出、打印等操作的前提是每行数据自足——当前文档的设计方式与此前提互斥。'
)

# ── 改造方案 ──
H1('五、改造方案：拆表')

P('改造思路：将原单一工作表拆分为三张独立表，每张仅承担一种职能，行粒度统一，列仅保留该职能所需。')

H2('表一：材料主数据')
P(
    '全部可选材料的属性目录。每行一个 SKU，所有列填满（废除隐式继承），'
    '增设「状态」列（启用 / 停用）替代 G 列的自由文本标注。保留约 12 列（原 25 列），'
    '删除公式文本和产品关联。排序、筛选、导出均安全可用。'
)

H2('表二：产品物料清单（BOM）')
P(
    '以树形结构定义成品与材料的组成关系。平板木饰面 → E0 多层板 + 木皮 + 封闭漆 + 封边条 + 辅材 + 加工费。'
    '层级列标识成品 / 构件 / 原材料，材料编码引用表一。层板的骨架材料作为独立行列出，'
    '不再隐藏于备注。空白成本项（门套线、门芯料、五金等）预先占位。'
)

H2('表三：全局参数')
P(
    '集中存放跨材料共享的定义：各材料大类损耗率、封边条统一单价、辅材和加工费综合单价、标准板材规格、'
    '投影面积定义（从当前 6 处重复收敛为 1 处）。修改一处即全局生效。'
)

H2('日常入口：本项目使用清单')
P(
    '从表一筛选状态 =「启用」，即获得仅含本项目材料的约 7 行清单。日常报价、核对、分享仅面对这 7 行。'
    '查阅完整目录或修改 BOM 时再切换至对应表——低频操作不应拖累高频操作。'
)

# ── 对比 ──
H1('六、改造前后对比')

add_simple_table(
    ['场景', '改造前', '改造后'],
    [
        ['查看本项目用了哪些材料', '56 行中逐行跳过不使用行，人工汇总', '筛选状态 = 启用，瞬间 7 行'],
        ['按单价排序', '排序后空行脱离上下文，数据身份丢失', '每行自足，排序安全'],
        ['按厚度筛选', '结果全部无名称无等级，不可用', '每行完整，筛选结果直接可用'],
        ['修改投影面积定义', '找到并修改 6 个单元格，确保一致', '改 1 个单元格'],
        ['导出给 ERP', '需向上填充脚本预处理 + 人工验证', '直接导出'],
        ['新人理解结构', '需逆向归纳 25 列的隐含规则', '3 张表，每张职能和列定义清晰'],
    ],
)

# ── 实施 ──
H1('七、实施建议')

B('第一轮（解决每次打开都遇到的问题）：')
B('- 材料主数据独立成表，每行填满，设状态列，废除隐式继承')
B('- 建立「本项目使用」筛选视图，活跃数据与参考数据物理分离')
B('- 公式文本移至全局参数表，每个定义只存一份')

P('')

B('第二轮（提升结构完整性）：')
B('- 产品 BOM 独立成表，树形表达成品材料组成，层板子项显式化')
B('- 补全「待确定」项判定，清理空行空列，补充版本号和修订记录')

# ── 保存 ──
doc.save(OUTPUT)
print('Done: ' + OUTPUT)
